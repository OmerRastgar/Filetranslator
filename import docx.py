import os
import docx
from deep_translator import GoogleTranslator

# Function to read the second column from a table in a docx file
def read_second_column(docx_file):
    doc = docx.Document(docx_file)
    tables = doc.tables
    return tables

# Function to translate a list of texts using GoogleTranslator
def translate_texts(texts, dest_language='es'):
    translator = GoogleTranslator(target=dest_language)
    translated_texts = [translator.translate(text) for text in texts]
    return translated_texts

# Function to write the translated texts into a new docx file
def write_translated_column_to_docx(tables, translated_texts, specific_phrase_translated, output_file):
    doc = docx.Document()
    
    for table_index, table in enumerate(tables):
        # Create a new table with the same number of rows and one column for the translation
        new_table = doc.add_table(rows=len(table.rows), cols=1)
        
        for row_index, row in enumerate(table.rows):
            cell = row.cells[1]  # Second column
            new_cell = new_table.cell(row_index, 0)
            
            # Original text for formatting preservation
            original_text = cell.text
            
            # Replace the text in the new cell with the translated text
            if row_index < len(translated_texts):
                new_cell.text = ''
                translated_text = translated_texts[row_index]
                
                # Detect the translated version of the specific phrase
                if specific_phrase_translated in translated_text:
                    add_formatted_phrase(new_cell, translated_text)
                else:
                    new_paragraph = new_cell.add_paragraph()
                    new_paragraph.add_run(translated_text).bold = True

    doc.save(output_file)

def add_formatted_phrase(cell, translated_text):
    phrases = translated_text.split('\n')
    
    # Adding the heading
    heading = cell.add_paragraph()
    heading_run = heading.add_run(phrases[0])
    heading_run.bold = True
    heading_run.underline = True
    
    # Adding bullet points
    for phrase in phrases[1:]:
        bullet_paragraph = cell.add_paragraph(style='List Bullet')
        bullet_paragraph.add_run(phrase.strip())

# Main program
input_docx_file = os.path.abspath('C:/Users/Omer Rastgar/Downloads/transalate/input.docx')
output_docx_file = os.path.abspath('translated_output.docx')
destination_language = 'es'  # Spanish

# Translate the specific phrase
specific_phrase = "ga has verified:"
translator = GoogleTranslator(target=destination_language)
specific_phrase_translated = translator.translate(specific_phrase)

# Read second column data
tables = read_second_column(input_docx_file)

# Collect texts from the second column of all tables
second_column_texts = []
for table in tables:
    for row in table.rows:
        second_column_texts.append(row.cells[1].text)

# Translate the texts
translated_texts = translate_texts(second_column_texts, dest_language=destination_language)

# Write translated texts to a new docx file while maintaining the original table style
write_translated_column_to_docx(tables, translated_texts, specific_phrase_translated, output_docx_file)

print(f'Translation complete. The translated data is saved in {output_docx_file}')
