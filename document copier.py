from docx import Document
from deep_translator import GoogleTranslator


def extract_user_data(template_path, filled_path):
    template_doc = Document(template_path)
    filled_doc = Document(filled_path)
    
    template_texts = [p.text for p in template_doc.paragraphs]
    user_data = {}
    translator = GoogleTranslator(target='es')
    print("starting extracting")
    # Extract paragraphs
    for i, p in enumerate(filled_doc.paragraphs):
        if i >= len(template_texts) or p.text != template_texts[i]:
            user_data[f"para_{i}"] = p.text
    print("extracted txt")
    # Extract table data
    
    for table_index, table in enumerate(filled_doc.tables):
        if table_index >= len(template_doc.tables):
            continue
        template_table = template_doc.tables[table_index]
        for row_index, row in enumerate(table.rows):
            if row_index >= len(template_table.rows):
                continue
            template_row = template_table.rows[row_index]
            for col_index, cell in enumerate(row.cells):
                if col_index >= len(template_row.cells):
                    continue
                template_cell_text = template_row.cells[col_index].text
                if cell.text != template_cell_text:
                    print(f"table No. {table_index} is extracting")
                    if (table_index == 4):
                        user_data[f"table_{table_index}_row_{row_index}_col_{col_index}"] = translator.translate(cell.text)
                    else:
                        user_data[f"table_{table_index}_row_{row_index}_col_{col_index}"] = cell.text
    print("extracted table")
    return user_data

def fill_target_language_template(target_template_path, user_data, output_path):
    target_doc = Document(target_template_path)
    print("adding text")
    # Insert paragraphs
    for key, value in user_data.items():
        if key.startswith("para_"):
            para_index = int(key.split("_")[1])
            if para_index < len(target_doc.paragraphs):
                target_doc.paragraphs[para_index].text = value
    print("adding table")
    # Insert table data
    for key, value in user_data.items():
        if key.startswith("table_"):
            parts = key.split("_")
            table_index = int(parts[1])
            row_index = int(parts[3])
            col_index = int(parts[5])
            if table_index < len(target_doc.tables):
                table = target_doc.tables[table_index]
                if row_index < len(table.rows):
                    row = table.rows[row_index]
                    if col_index < len(row.cells):
                        row.cells[col_index].text = value
    
    target_doc.save(output_path)

# Paths to the documents
template_path = "template_english.docx"
filled_path = "GURBP-CO-240425-01-SA-R1  - Rubyplay Security Report (EN) V1.0.docx"
target_template_path = "ClientID-CO-YYMMDD-XX-SA-R1  - ClientName Security Report (ES) V0.1.docx"
output_path = "filled_target_language.docx"

# Extract user-filled data from the English filled report
user_data = extract_user_data(template_path, filled_path)

# Fill the target language template with the user-filled data
fill_target_language_template(target_template_path, user_data, output_path)



